"""
DOC (Legacy Word) text extraction using python-docx with fallback
Note: python-docx primarily supports .docx files. For .doc files, 
the library will attempt to read them, but may have limited success.
For full .doc support, consider using additional libraries like:
- antiword (requires external binary)
- pywin32 (Windows only, requires Microsoft Word installed)
- textract (requires additional dependencies)

For now, we'll attempt to use python-docx for both formats.
"""
from docx import Document
from pathlib import Path
from typing import Dict


class DOCExtractor:
    """
    Extract text from legacy DOC files
    Note: python-docx has limited .doc support. May work for some files.
    For better .doc support, consider converting to .docx first or using antiword.
    """
    
    def __init__(self):
        self.supported_extension = '.doc'
    
    def extract_text(self, file_path: str) -> Dict[str, any]:
        """
        Attempt to extract text from a DOC file using python-docx
        
        Args:
            file_path: Path to the DOC file
            
        Returns:
            Dictionary containing:
                - text: Extracted text
                - paragraphs: Number of paragraphs
                - filename: Original filename
                - warning: Any warnings about extraction
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.suffix.lower() != self.supported_extension:
            raise ValueError(f"Invalid file type. Expected {self.supported_extension}, got {path.suffix}")
        
        try:
            # Attempt to open with python-docx
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n" + "\n".join(table_texts)
            
            result = {
                'text': all_text,
                'paragraphs': len(paragraphs),
                'filename': path.name,
                'file_type': 'doc'
            }
            
            # Add warning if text extraction seems incomplete
            if not all_text.strip():
                result['warning'] = "⚠️ Text extraction may be incomplete. Legacy .doc format has limited support. Consider converting to .docx for better results."
            
            return result
            
        except Exception as e:
            # If python-docx fails, provide helpful error message
            raise Exception(
                f"Error reading .doc file: {str(e)}\n\n"
                "💡 Suggestion: Legacy .doc files have limited support. "
                "Please convert to .docx format for better results. "
                "You can convert using Microsoft Word or online converters."
            )
