"""
DOCX text extraction using python-docx
"""
from docx import Document
from pathlib import Path
from typing import Dict


class DOCXExtractor:
    """Extract text from DOCX files"""
    
    def __init__(self):
        self.supported_extension = '.docx'
    
    def extract_text(self, file_path: str) -> Dict[str, any]:
        """
        Extract text from a DOCX file
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary containing:
                - text: Extracted text
                - paragraphs: Number of paragraphs
                - filename: Original filename
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if path.suffix.lower() != self.supported_extension:
            raise ValueError(f"Invalid file type. Expected {self.supported_extension}, got {path.suffix}")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Only include non-empty paragraphs
                    paragraphs.append(para.text)
            
            # Extract text from tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        table_texts.append(" | ".join(row_text))
            
            # Combine all text
            all_text = "\n\n".join(paragraphs)
            if table_texts:
                all_text += "\n\n" + "\n".join(table_texts)
            
            return {
                'text': all_text,
                'paragraphs': len(paragraphs),
                'filename': path.name,
                'file_type': 'docx'
            }
            
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}")
